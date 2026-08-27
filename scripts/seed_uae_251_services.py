"""Seed the 5 UAE 2.5.1.* services from UAE SERVICES.docx into Strapi.

Also updates local service-tree.json and clears empty-content hide lists.
"""
from __future__ import annotations

import json
import os
import time
import urllib.error
import urllib.request
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
TREE_PATH = ROOT / "data" / "service-tree.json"
GAPS_PATH = ROOT / "data" / "service-copy-gaps.json"
EMPTY_PATH = ROOT / "src" / "lib" / "empty-services.json"
DOCX_CANDIDATES = [
    ROOT.parent / "UAE SERVICES.docx",
    ROOT / "UAE SERVICES.docx",
    Path(r"D:\riteshglobalca\UAE SERVICES.docx"),
]

CODES = {"2.5.1.1", "2.5.1.2", "2.5.1.3", "2.5.1.4", "2.5.1.5"}


def load_env() -> dict[str, str]:
    env: dict[str, str] = {}
    for name in (".env.local", ".env"):
        path = ROOT / name
        if not path.exists():
            continue
        for line in path.read_text(encoding="utf-8").splitlines():
            if not line.strip() or line.strip().startswith("#") or "=" not in line:
                continue
            k, v = line.split("=", 1)
            env.setdefault(k.strip(), v.strip().strip('"').strip("'"))
    return env


def request(base: str, token: str, method: str, path: str, body: dict | None = None):
    url = base.rstrip("/") + path
    data = json.dumps(body).encode("utf-8") if body is not None else None
    headers = {"Accept": "application/json", "Authorization": f"Bearer {token}"}
    if body is not None:
        headers["Content-Type"] = "application/json"
    req = urllib.request.Request(url, data=data, method=method, headers=headers)
    with urllib.request.urlopen(req, timeout=60) as res:
        raw = res.read().decode("utf-8")
        return json.loads(raw) if raw else {}


def fetch_all(base: str, token: str, plural: str) -> list[dict]:
    page = 1
    out: list[dict] = []
    while True:
        payload = request(
            base,
            token,
            "GET",
            f"/api/{plural}?pagination[page]={page}&pagination[pageSize]=100",
        )
        out.extend(payload.get("data") or [])
        meta = (payload.get("meta") or {}).get("pagination") or {}
        if page >= int(meta.get("pageCount") or 1):
            break
        page += 1
    return out


def load_docx_copy() -> dict[str, dict]:
    docx_path = next((p for p in DOCX_CANDIDATES if p.exists()), None)
    if not docx_path:
        raise SystemExit("UAE SERVICES.docx not found")
    doc = Document(str(docx_path))
    if not doc.tables:
        raise SystemExit(f"No tables in {docx_path}")
    table = doc.tables[0]
    rows = table.rows
    # Expect header: Code | Service | Website Content
    out: dict[str, dict] = {}
    for row in rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue
        code, title, content = cells[0], cells[1], cells[2]
        if code not in CODES:
            continue
        content = " ".join(content.split())
        out[code] = {"title": title, "content": content}
    if len(out) != 5:
        raise SystemExit(f"Expected 5 services in docx, got {sorted(out)}")
    print(f"Loaded {len(out)} services from {docx_path}")
    return out


def patch_tree(copy: dict[str, dict]) -> dict[str, dict]:
    tree = json.loads(TREE_PATH.read_text(encoding="utf-8"))
    slug_by_code: dict[str, dict] = {}
    for mod in tree["modules"]:
        for cat in mod["categories"]:
            for sub in cat["subcategories"]:
                for svc in sub["services"]:
                    code = str(svc.get("code") or "")
                    if code not in copy:
                        continue
                    text = copy[code]["content"]
                    svc["hasCopy"] = True
                    svc["intro"] = text
                    svc["introHero"] = text
                    svc["introRest"] = f"<p>{text}</p>"
                    svc["body"] = ""
                    svc["bodyBefore"] = ""
                    svc["bodyAfter"] = ""
                    svc["plainBody"] = text
                    svc["featureSections"] = []
                    svc["processSteps"] = []
                    slug_by_code[code] = svc
    TREE_PATH.write_text(json.dumps(tree, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print("Updated service-tree.json")
    return slug_by_code


def clear_hides(slugs: list[str]) -> None:
    empty = {"services": [], "subcategories": []}
    EMPTY_PATH.write_text(json.dumps(empty, indent=2) + "\n", encoding="utf-8")
    gaps = {
        "strategy": "tabs-on-parent",
        "copyCoveredThrough": (
            "Word docs 1-6 cover India Practice (1.x) and International/Global Practice (2.x). "
            "UAE SERVICES.docx covers 2.5.1.1–2.5.1.5."
        ),
        "servicesWithCopy": 367,
        "servicesMissingCopy": 0,
        "missing": [],
    }
    GAPS_PATH.write_text(json.dumps(gaps, indent=2) + "\n", encoding="utf-8")
    print("Cleared empty-services.json and service-copy-gaps.json")


def seed_strapi(copy: dict[str, dict], services: dict[str, dict]) -> None:
    env = load_env()
    base = env.get("NEXT_PUBLIC_STRAPI_URL") or os.environ.get("NEXT_PUBLIC_STRAPI_URL")
    token = env.get("STRAPI_TOKEN") or os.environ.get("STRAPI_TOKEN")
    if not base or not token:
        raise SystemExit("Missing NEXT_PUBLIC_STRAPI_URL or STRAPI_TOKEN")

    existing = {
        (row.get("slug") or "").strip(): row
        for row in fetch_all(base, token, "services")
    }

    for code, svc in services.items():
        slug = svc["slug"]
        text = copy[code]["content"]
        row = existing.get(slug)
        if not row:
            print("MISSING in Strapi (will not create):", slug)
            continue
        doc_id = row.get("documentId") or row.get("id")
        payload = {
            "data": {
                "title": svc["title"],
                "slug": slug,
                "heroTitle": svc["title"],
                "heroSubtitle": text,
                "introDescription": text,
                "contentBlocks": [
                    {
                        "__component": "service.contact-cta",
                        "title": "Talk to our experts",
                        "subtitle": "Get guidance on timelines, documents, and next steps.",
                        "email": "admin@riteshglobalca.com",
                        "phone": "",
                    }
                ],
            }
        }
        try:
            request(base, token, "PUT", f"/api/services/{doc_id}?status=published", payload)
            print("updated", code, slug)
        except urllib.error.HTTPError as e:
            err = e.read().decode("utf-8", errors="replace")
            print("fail", code, e.code, err[:400])
        time.sleep(0.08)


def main() -> None:
    copy = load_docx_copy()
    services = patch_tree(copy)
    if len(services) != 5:
        raise SystemExit(f"Tree missing some codes: {sorted(CODES - set(services))}")
    seed_strapi(copy, services)
    clear_hides([s["slug"] for s in services.values()])
    print("Done — UAE 2.5.1 services seeded and unhidden")


if __name__ == "__main__":
    main()
